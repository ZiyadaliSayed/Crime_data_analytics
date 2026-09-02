from docx import Document
from docx.shared import Pt, RGBColor
import re

def create_docx_from_md(md_file, docx_file):
    doc = Document()
    
    with open(md_file, 'r', encoding='utf-8') as f:
        lines = f.readlines()
        
    in_code_block = False
    
    for line in lines:
        line = line.strip('\n')
        
        if line.startswith('```'):
            in_code_block = not in_code_block
            continue
            
        if in_code_block:
            p = doc.add_paragraph(line)
            p.style = 'No Spacing'
            if p.runs:
                p.runs[0].font.name = 'Courier New'
            continue
            
        if not line.strip():
            continue
            
        if line.startswith('## '):
            doc.add_heading(line[3:], level=2)
        elif line.startswith('### '):
            doc.add_heading(line[4:], level=3)
        elif line.startswith('# '):
            doc.add_heading(line[2:], level=1)
        elif line.startswith('*   ') or line.startswith('- '):
            p = doc.add_paragraph(style='List Bullet')
            text = line.strip('* -')
            add_formatted_text(p, text)
        elif re.match(r'^\d+\.\s', line):
            p = doc.add_paragraph(style='List Number')
            text = re.sub(r'^\d+\.\s', '', line)
            add_formatted_text(p, text)
        elif line.startswith('---'):
            doc.add_page_break()
        else:
            p = doc.add_paragraph()
            add_formatted_text(p, line)
            
    doc.save(docx_file)

def add_formatted_text(paragraph, text):
    # Extremely basic bold processing
    parts = text.split('**')
    for i, part in enumerate(parts):
        run = paragraph.add_run(part)
        if i % 2 != 0:
            run.bold = True

if __name__ == "__main__":
    create_docx_from_md('Project_Report.md', 'Project_Report.docx')
    print("Successfully generated Project_Report.docx")
