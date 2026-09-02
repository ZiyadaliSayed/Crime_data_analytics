from docx import Document
from docx.shared import Pt

doc = Document()
doc.add_heading('Crime Data Analytics Project - Dataset Information', 0)

doc.add_heading('1. Real Crime Data Scraped (2023/2024 proxy)', level=2)
doc.add_paragraph('Source Link: https://ncrb.gov.in/en/crime-in-india')
doc.add_paragraph('Usage: This dataset provides the highly detailed, state-wise breakdown of various crimes (Murder, Rape, Kidnapping, Robbery, etc.). Because the raw 2024 crime dataset is not yet publicly accessible in a structured format, this 2023 dataset is used as a 1:1 proxy for 2024 to perfectly align with the 2024 literacy survey for correlative analysis.')
doc.add_paragraph('Columns Used:')
doc.add_paragraph('• State / UT', style='List Bullet')
doc.add_paragraph('• Total Crimes (IPC+SLL) 2023', style='List Bullet')
doc.add_paragraph('• Crime Rate (IPC+SLL) 2023', style='List Bullet')
doc.add_paragraph('• Murder 2023, Rape 2023, Kidnapping 2023, Robbery & Dacoity 2023, Hit & Run 2023, Illegal arms 2023, Corruption (Total cases) 2023', style='List Bullet')

doc.add_heading('2. Indian Cities Demographics', level=2)
doc.add_paragraph('Source Link: https://censusindia.gov.in/ (or Open Government Data via Kaggle)')
doc.add_paragraph('Usage: This dataset is used to extract the total urban populations for each state. This population data is mathematically necessary to accurately calculate the true "Crime Rate per 1 Lakh people" rather than just relying on raw, unadjusted crime totals.')
doc.add_paragraph('Columns Used:')
doc.add_paragraph('• state_name', style='List Bullet')
doc.add_paragraph('• population_total (summed by state to calculate Total_Urban_Population)', style='List Bullet')
doc.add_paragraph('• effective_literacy_rate_total (used as a fallback average)', style='List Bullet')

doc.add_heading('3. Indian Prison Statistics', level=2)
doc.add_paragraph('Source Link: https://ncrb.gov.in/en/prison-statistics-india')
doc.add_paragraph('Usage: This dataset is used to build the Dim_Prison_Stats dimension table. It provides contextual socio-economic data regarding the prison population, specifically focusing on the educational backgrounds of inmates to support the illiteracy correlation.')
doc.add_paragraph('Columns Used:')
doc.add_paragraph('• State/UT', style='List Bullet')
doc.add_paragraph('• Total - Total (mapped to Total_Prisoners)', style='List Bullet')
doc.add_paragraph('• Educational Standard - Illiterate (mapped to Illiterate_Prisoners)', style='List Bullet')
doc.add_paragraph('• Educational Standard - Graduate (mapped to Graduate_Prisoners)', style='List Bullet')

doc.add_heading('4. Wikipedia: Crime in India (Live Web Scrape)', level=2)
doc.add_paragraph('Source Link: https://en.wikipedia.org/wiki/Crime_in_India')
doc.add_paragraph('Usage: The ETL pipeline live-scrapes the HTML tables on this page to extract the historical total crime records registered in 2017 for every state. This acts as the pure data baseline for the exact 2017 to 2024 comparisons.')
doc.add_paragraph('Columns Used:')
doc.add_paragraph('• State/UT', style='List Bullet')
doc.add_paragraph('• 2017 (Historical total crime count)', style='List Bullet')

doc.add_heading('5. Wikipedia: Indian States by Literacy Rate (Live Web Scrape)', level=2)
doc.add_paragraph('Source Link: https://en.wikipedia.org/wiki/List_of_Indian_states_and_union_territories_by_literacy_rate')
doc.add_paragraph('Usage: The pipeline actively scrapes this page to extract the official 2017 National Statistical Office (NSO) Survey and the 2023-2024 Periodic Labour Force Survey (PLFS). This provides the exact literacy and illiteracy percentages used to prove the mathematical correlation between illiteracy and crime.')
doc.add_paragraph('Columns Used:')
doc.add_paragraph('• State or UT (Column 0)', style='List Bullet')
doc.add_paragraph('• 2017 Total (Column 4 - NSO Survey)', style='List Bullet')
doc.add_paragraph('• 2024 Total (Column 7 - PLFS Survey)', style='List Bullet')

doc.save('Dataset_Information.docx')
