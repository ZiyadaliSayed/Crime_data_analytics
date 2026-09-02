from docx import Document
from docx.shared import Pt, Inches

doc = Document()

# Add a bold, center-aligned title
title = doc.add_heading('SAMPLE SOLUTION OF PROBLEM STMT1', 0)
title.alignment = 1

doc.add_paragraph('Project Title\nCrime Data Warehouse for Socio-Economic Analytics', style='Body Text').runs[0].bold = True
doc.add_paragraph('Objective\nDesign and implement a data warehouse by integrating crime statistics and demographic data from multiple sources. Build dimensional models, perform ETL, create OLAP cubes, and visualize business insights using dashboards to study the correlation between crime and illiteracy.', style='Body Text').runs[0].bold = True

doc.add_paragraph('_' * 80)

# Phase 1
doc.add_paragraph('Phase 1: Requirement Analysis', style='Body Text').runs[0].bold = True
doc.add_paragraph('Step 1: Study Business Requirements\nIdentify the business questions.\nExamples:')
doc.add_paragraph('• Which state generates the highest crime rate?', style='List Bullet')
doc.add_paragraph('• Which state has the highest illiteracy rate?', style='List Bullet')
doc.add_paragraph('• What is the mathematical correlation between illiteracy and crime?', style='List Bullet')
doc.add_paragraph('• How has national crime evolved over time (2017 vs 2024)?', style='List Bullet')
doc.add_paragraph('• What is the breakdown of violent crimes across states?', style='List Bullet')
doc.add_paragraph('Deliverable\nBusiness Requirement Document', style='Body Text').runs[0].bold = True

doc.add_paragraph('_' * 80)

# Phase 2
doc.add_paragraph('Phase 2: Data Collection', style='Body Text').runs[0].bold = True
doc.add_paragraph('Collect data from multiple sources.\nExample Sources')
doc.add_paragraph('Source 1: crime_dataset_india.csv (40,000+ incidents Kaggle Dataset)\n| City | Date of Occurrence | Crime Description | Report Number |')
doc.add_paragraph('Source 2: ncrb_2024_statewise_supplement.csv (Used to supplement missing states in the Kaggle sample)\n| State / UT | Total Crimes | Crime Rate | Murder | Rape | Kidnapping | Extortion |')
doc.add_paragraph('Source 3: indian_cities_demographics.csv\n| state_name | population_total | effective_literacy_rate_total |')
doc.add_paragraph('Source 4: wikipedia_crime_in_india.csv (Local CSV)\n| State/UT | 2017 |')
doc.add_paragraph('Source 5: wikipedia_literacy_rates.csv (Local CSV)\n| State or UT | 2017 Total | 2024 Total |')
doc.add_paragraph('Deliverable\nRaw datasets', style='Body Text').runs[0].bold = True

doc.add_paragraph('_' * 80)

# Phase 3
doc.add_paragraph('Phase 3: Data Preprocessing', style='Body Text').runs[0].bold = True
doc.add_paragraph('Perform')
doc.add_paragraph('• Remove missing values from scraped tables', style='List Bullet')
doc.add_paragraph('• Standardize state names across all datasets', style='List Bullet')
doc.add_paragraph('• Convert string numbers with commas to integers', style='List Bullet')
doc.add_paragraph('• Calculate missing demographic fields (e.g., Illiteracy_Rate)', style='List Bullet')
doc.add_paragraph('Example\n"Nct Of Delhi"\n"Delhi"\n↓\nDelhi')
doc.add_paragraph('Deliverable\nClean datasets', style='Body Text').runs[0].bold = True

doc.add_paragraph('_' * 80)

# Phase 4
doc.add_paragraph('Phase 4: ETL Pipeline', style='Body Text').runs[0].bold = True
doc.add_paragraph('Extract\nRead CSV files\ncrime_dataset_india.csv\nncrb_2024_statewise_supplement.csv\nindian_cities_demographics.csv\nwikipedia_crime_in_india.csv\nwikipedia_literacy_rates.csv')
doc.add_paragraph('Transform\nPerform')
doc.add_paragraph('• Merge tables based on State_Name', style='List Bullet')
doc.add_paragraph('• Data cleaning', style='List Bullet')
doc.add_paragraph('• Surrogate key generation', style='List Bullet')
doc.add_paragraph('• Create derived columns', style='List Bullet')
doc.add_paragraph('Example\nCrime Rate per 1 Lakh\n= (Total_Crimes / Total_Urban_Population) * 100000')
doc.add_paragraph('Load\nLoad into\nSQLite')
doc.add_paragraph('Deliverable\nWorking ETL Pipeline', style='Body Text').runs[0].bold = True
doc.add_paragraph('Tools')
doc.add_paragraph('• Python (Pandas, Requests)', style='List Bullet')
doc.add_paragraph('• SQLite3', style='List Bullet')
doc.add_paragraph('• Streamlit', style='List Bullet')

doc.add_paragraph('_' * 80)

# Phase 5
doc.add_paragraph('Phase 5: Design Star Schema', style='Body Text').runs[0].bold = True
doc.add_paragraph('Fact Table\nFact_Crime_Stats\n| State_ID | Year | Total_Crimes | Crime_Rate | Literacy_Rate | Murder | Rape | Kidnapping | Extortion | Robbery_Dacoity | Hit_Run | Illegal_Arms | Corruption |')
doc.add_paragraph('Dimension Tables\nDim_State\n| State_ID | State_Name | Total_Urban_Population | Avg_Literacy_Rate |')

doc.add_paragraph('Deliverable\nER Diagram', style='Body Text').runs[0].bold = True

doc.add_paragraph('_' * 80)

# Phase 6
doc.add_paragraph('Phase 6: Design Snowflake Schema', style='Body Text').runs[0].bold = True
doc.add_paragraph('Normalize dimensions.\nExample\nDim_State\n↓\nRegion Table\n↓\nCountry Table')
doc.add_paragraph('Deliverable\nSnowflake Schema Diagram', style='Body Text').runs[0].bold = True

doc.add_paragraph('_' * 80)

# Phase 7
doc.add_paragraph('Phase 7: Create Data Warehouse', style='Body Text').runs[0].bold = True
doc.add_paragraph('Create SQL tables\nCREATE TABLE Dim_State\nCREATE TABLE Fact_Crime_Stats\nLoad transformed data.')
doc.add_paragraph('Deliverable\nOperational Data Warehouse', style='Body Text').runs[0].bold = True

doc.add_paragraph('_' * 80)

# Phase 8
doc.add_paragraph('Phase 8: Create Data Cube', style='Body Text').runs[0].bold = True
doc.add_paragraph('Dimensions')
doc.add_paragraph('• State', style='List Bullet')
doc.add_paragraph('• Year', style='List Bullet')
doc.add_paragraph('• Crime Type', style='List Bullet')
doc.add_paragraph('Measures')
doc.add_paragraph('• Total Crimes', style='List Bullet')
doc.add_paragraph('• Crime Rate', style='List Bullet')
doc.add_paragraph('• Illiteracy Rate', style='List Bullet')
doc.add_paragraph('Cube\n        Time\n         |\n         |\nLocation ---- Crime Counts\n         |\n         |\n    Crime Type')
doc.add_paragraph('Deliverable\nData Cube', style='Body Text').runs[0].bold = True

doc.add_paragraph('_' * 80)

# Phase 9
doc.add_paragraph('Phase 9: Perform OLAP Operations', style='Body Text').runs[0].bold = True
doc.add_paragraph('Roll-up\nState Crimes\n↓\nNational Crimes\nExample SQL\nGROUP BY Year')
doc.add_paragraph('Drill-down\nNational Crime\n↓\nState Crime\n↓\nSpecific Violent Crime (Murder, Rape)')
doc.add_paragraph('Slice\nOnly\nYear = 2024')
doc.add_paragraph('Dice\nYear = 2024\nState = Maharashtra')
doc.add_paragraph('Pivot\nRows (States)\n↓\nColumns (Years)\nInterchange dimensions')
doc.add_paragraph('Deliverable\nScreenshots of each OLAP operation', style='Body Text').runs[0].bold = True

doc.add_paragraph('_' * 80)

# Phase 10
doc.add_paragraph('Phase 10: Business Intelligence Dashboard', style='Body Text').runs[0].bold = True
doc.add_paragraph('Use\nStreamlit\nPlotly Visualizations')
doc.add_paragraph('KPI Cards')
doc.add_paragraph('• National Total Crimes', style='List Bullet')
doc.add_paragraph('• National Crime Rate', style='List Bullet')
doc.add_paragraph('• Avg National Illiteracy (%)', style='List Bullet')
doc.add_paragraph('Charts\nState Crime Ranking (Horizontal Bar Chart)\nState Illiteracy Ranking (Horizontal Bar Chart)\nIlliteracy vs Crime Correlation (Scatter Plot with Trendline)\nViolent Crime Breakdown (Grouped Bar Chart)\nNational Crime Evolution (Line Chart)\nState Comparison 2017 vs 2024 (Grouped Bar Charts)')
doc.add_paragraph('Interactive Filters')
doc.add_paragraph('• Year Selection', style='List Bullet')
doc.add_paragraph('• Multi-State Comparison Filter', style='List Bullet')
doc.add_paragraph('• Violent Crime Type Filter', style='List Bullet')
doc.add_paragraph('• Logarithmic Scale Toggle', style='List Bullet')
doc.add_paragraph('Deliverable\nInteractive Dashboard', style='Body Text').runs[0].bold = True

doc.add_paragraph('_' * 80)

# Phase 11
doc.add_paragraph('Phase 11: Testing', style='Body Text').runs[0].bold = True
doc.add_paragraph('Validate')
doc.add_paragraph('• ETL accuracy', style='List Bullet')
doc.add_paragraph('• Duplicate removal', style='List Bullet')
doc.add_paragraph('• Correct aggregation', style='List Bullet')
doc.add_paragraph('• Dashboard correctness', style='List Bullet')
doc.add_paragraph('• Query execution', style='List Bullet')
doc.add_paragraph('Deliverable\nTesting Report', style='Body Text').runs[0].bold = True

doc.add_paragraph('_' * 80)

# Phase 12
doc.add_paragraph('Phase 12: Documentation', style='Body Text').runs[0].bold = True
doc.add_paragraph('Include')
doc.add_paragraph('• Problem Statement', style='List Bullet')
doc.add_paragraph('• Objectives', style='List Bullet')
doc.add_paragraph('• Literature Review', style='List Bullet')
doc.add_paragraph('• System Architecture', style='List Bullet')
doc.add_paragraph('• Database Design', style='List Bullet')
doc.add_paragraph('• ETL Process', style='List Bullet')
doc.add_paragraph('• Star Schema', style='List Bullet')
doc.add_paragraph('• Snowflake Schema', style='List Bullet')
doc.add_paragraph('• SQL Queries', style='List Bullet')
doc.add_paragraph('• OLAP Operations', style='List Bullet')
doc.add_paragraph('• Dashboard Screenshots', style='List Bullet')
doc.add_paragraph('• Results', style='List Bullet')
doc.add_paragraph('• Future Scope', style='List Bullet')

doc.add_paragraph('_' * 80)

# Software Requirements
doc.add_paragraph('Software Requirements', style='Body Text').runs[0].bold = True
table = doc.add_table(rows=6, cols=2)
table.cell(0, 0).text = 'Component'
table.cell(0, 1).text = 'Tool'
table.cell(1, 0).text = 'Programming'
table.cell(1, 1).text = 'Python'
table.cell(2, 0).text = 'Database'
table.cell(2, 1).text = 'SQLite3'
table.cell(3, 0).text = 'ETL'
table.cell(3, 1).text = 'Pandas / Requests'
table.cell(4, 0).text = 'Visualization'
table.cell(4, 1).text = 'Streamlit / Plotly'
table.cell(5, 0).text = 'IDE'
table.cell(5, 1).text = 'VS Code'

for i in range(6):
    for j in range(2):
        if i == 0:
            table.cell(i, j).paragraphs[0].runs[0].bold = True

# Expected Outputs
doc.add_paragraph('\nExpected Outputs', style='Body Text').runs[0].bold = True
table2 = doc.add_table(rows=8, cols=2)
table2.cell(0, 0).text = 'Outcome'
table2.cell(0, 1).text = 'Description'
table2.cell(1, 0).text = 'ETL Pipeline'
table2.cell(1, 1).text = 'Extract, transform, and load data from local CSVs into the warehouse'
table2.cell(2, 0).text = 'Star Schema'
table2.cell(2, 1).text = 'Fact table linked to State dimension'
table2.cell(3, 0).text = 'Snowflake Schema'
table2.cell(3, 1).text = 'Normalized dimension tables (e.g., State hierarchy)'
table2.cell(4, 0).text = 'Data Warehouse'
table2.cell(4, 1).text = 'Centralized relational warehouse in SQLite'
table2.cell(5, 0).text = 'Data Cube'
table2.cell(5, 1).text = 'Multi-dimensional model with State, Time, and Crime Type dimensions'
table2.cell(6, 0).text = 'OLAP Operations'
table2.cell(6, 1).text = 'Demonstration of Roll-up, Drill-down, Slice, Dice, and Pivot using SQL'
table2.cell(7, 0).text = 'Interactive Dashboard'
table2.cell(7, 1).text = 'Streamlit dashboard with KPIs, charts, filters, and drill-through analysis'

for i in range(8):
    for j in range(2):
        if i == 0:
            table2.cell(i, j).paragraphs[0].runs[0].bold = True

# Suggested Timeline
doc.add_paragraph('\nSuggested Timeline (4 Weeks)', style='Body Text').runs[0].bold = True
table3 = doc.add_table(rows=5, cols=2)
table3.cell(0, 0).text = 'Week'
table3.cell(0, 1).text = 'Activities'
table3.cell(1, 0).text = 'Week 1'
table3.cell(1, 1).text = 'Requirement analysis, data collection, preprocessing, and ETL development'
table3.cell(2, 0).text = 'Week 2'
table3.cell(2, 1).text = 'Design and implement Star/Snowflake schemas, create the data warehouse, and load data'
table3.cell(3, 0).text = 'Week 3'
table3.cell(3, 1).text = 'Build the data cube, perform OLAP operations, and validate analytical queries'
table3.cell(4, 0).text = 'Week 4'
table3.cell(4, 1).text = 'Develop the Streamlit dashboard, conduct testing, and complete documentation'

for i in range(5):
    for j in range(2):
        if i == 0:
            table3.cell(i, j).paragraphs[0].runs[0].bold = True

doc.add_paragraph('\nThis workflow exposes students to the complete Business Intelligence lifecycle, from raw operational data to dimensional modeling, ETL, OLAP analysis, and interactive decision-support dashboards, making it well suited as a comprehensive undergraduate mini-project.')

doc.save('Final_Project_Report_Formatted.docx')
